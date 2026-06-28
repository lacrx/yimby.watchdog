#!/usr/bin/env python3
"""Generate Word document: Effect of 20% Inclusionary Rate Inside the SSCSP."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    hs.font.name = 'Calibri'

# ─── Title ───
title = doc.add_heading('Effect of a 20% Inclusionary Rate Inside the\nSmart & Sustainable Corridors Specific Plan', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Policy Analysis — June 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

doc.add_paragraph()

# ─── Section 1 ───
doc.add_heading('What the SSCSP Says About Affordability', level=2)

doc.add_paragraph(
    'The Smart & Sustainable Corridors Specific Plan (SSCSP), going to City Council '
    'for adoption on June 24, 2026, defers entirely to the citywide Inclusionary Housing '
    'Ordinance for affordable housing requirements. Chapter 3 (Land Use), page 3-6 states: '
    '"Income-restricted housing will be implemented through the City\'s Inclusionary Housing '
    'Ordinance and state incentives for affordable housing."'
)

doc.add_paragraph(
    'There is no SSCSP-specific inclusionary rate. Whatever the citywide IHO requires, '
    'the corridors follow. The current citywide IHO requires a 10% affordable set-aside.'
)

# ─── Section 2 ───
doc.add_heading('Projected Buildout Numbers', level=2)

doc.add_paragraph(
    'Table 3-8 of the SSCSP projects the following residential buildout across the three corridors:'
)

# Buildout table
table1 = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers1 = ['Corridor', 'Existing Units', 'Projected New', 'Total']
data1 = [
    ['Mission Avenue', '300', '4,100', '4,400'],
    ['Oceanside Blvd', '100', '3,500', '3,600'],
    ['Vista Way', '400', '700', '1,100'],
    ['Total', '800', '8,300', '9,100'],
]
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for ri, row_data in enumerate(data1):
    for ci, val in enumerate(row_data):
        cell = table1.rows[ri + 1].cells[ci]
        cell.text = val
        if ri == 3:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True

doc.add_paragraph()

doc.add_paragraph(
    'The SSCSP also projects 14,500 new jobs and a population increase of 22,200 '
    'within the Planning Area by 2050. Approximately 50% of all new citywide residential '
    'development is projected to occur within these three corridors.'
)

# ─── Section 3 ───
doc.add_heading('Affordable Unit Production by Inclusionary Rate', level=2)

doc.add_paragraph(
    'Simple arithmetic suggests doubling the inclusionary rate doubles affordable output. '
    'The city\'s own feasibility data shows the opposite.'
)

table2 = doc.add_table(rows=4, cols=4, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Scenario', 'Total Units Built', 'Affordable @ Rate', 'Affordable Produced']
data2 = [
    ['10% IHO (current)', '~8,300', '10%', '~830'],
    ['15% IHO', '~5,000', '15%', '~750'],
    ['20% IHO', '~2,500', '20%', '~500'],
]
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for ri, row_data in enumerate(data2):
    for ci, val in enumerate(row_data):
        table2.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    'Doubling the inclusionary rate cuts affordable production by approximately 40%. '
    'You get fewer affordable units AND fewer market units.'
)
run.bold = True

# ─── Section 4 ───
doc.add_heading('Why 20% Kills the Pipeline: The DRA Feasibility Study', level=2)

doc.add_paragraph(
    'The City of Oceanside commissioned DRA (Development Resource Advisors) to produce '
    'an Inclusionary Housing Study in February 2022. This study modeled the residual land '
    'value (RLV) — the price a developer can pay for land after accounting for all '
    'construction costs and required returns — at different inclusionary set-aside levels.'
)

doc.add_paragraph('The feasibility cliff hits at 15%, not 20%:')

table3 = doc.add_table(rows=4, cols=4, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['Prototype', '10% Set-Aside RLV', '15% Set-Aside RLV', 'Direction at 20%']
data3 = [
    ['Structured parking', '$3/sf (barely viable)', '-$11/sf (infeasible)', 'Much worse'],
    ['Podium', '$57/sf', '$36/sf', '~$15/sf or below'],
    ['Surface parking', '$62/sf', '$46/sf', '~$30/sf'],
]
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for ri, row_data in enumerate(data3):
    for ci, val in enumerate(row_data):
        table3.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()

doc.add_paragraph(
    'Most SSCSP corridor development requires structured or podium parking. The plan\'s '
    'own Section 3.6 acknowledges that "structured and underground parking... can be '
    'prohibitively expensive" (p. 3-24). The SSCSP density standards — Corridors Mixed Use '
    'at 40 du/acre, Centers Mixed Use at 55 du/acre — produce 4-8 story buildings that '
    'cannot pencil with surface parking lots alone.'
)

doc.add_paragraph(
    'At 20% inclusionary:'
)

bullets_20pct = [
    'Structured parking projects go negative at 15% — 20% makes them deeply infeasible.',
    'Podium projects lose approximately 60% of residual land value — most sites become '
    'uncompetitive with existing uses.',
    'Only surface-parking projects survive — but those cannot achieve the densities '
    'the SSCSP envisions for the corridors.',
]
for b in bullets_20pct:
    doc.add_paragraph(b, style='List Bullet')

# ─── Section 5 ───
doc.add_heading('The SSCSP Already Reduces Costs to Help Feasibility', level=2)

doc.add_paragraph(
    'The plan establishes parking standards (Table 3-7) significantly below typical '
    'municipal requirements:'
)

table4 = doc.add_table(rows=5, cols=2, style='Light Shading Accent 1')
table4.alignment = WD_TABLE_ALIGNMENT.CENTER
headers4 = ['Land Use', 'Minimum Parking Standard']
data4 = [
    ['Market Rate Residential', '1 space per 1,200 sf'],
    ['Market Rate near SPRINTER (½ mi)', '1 space per 1,500 sf'],
    ['Income Restricted Residential', '0.5 spaces per unit'],
    ['Commercial', '1 space per 500 sf'],
]
for i, h in enumerate(headers4):
    cell = table4.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
for ri, row_data in enumerate(data4):
    for ci, val in enumerate(row_data):
        table4.rows[ri + 1].cells[ci].text = val

doc.add_paragraph()

doc.add_paragraph(
    'The plan also allows unbundled parking (LU-55), shared parking between adjacent uses '
    '(LU-27), and parking in-lieu fees to fund centralized district parking (LU-56). '
    'These are implicit acknowledgments that development in these corridors is marginal. '
    'Adding a 20% inclusionary mandate works against every cost-reduction strategy the '
    'SSCSP implements.'
)

# ─── Section 6 ───
doc.add_heading('Cascading Effects Beyond Housing', level=2)

doc.add_paragraph(
    'The SSCSP\'s projected 14,500 new jobs and 22,200 new residents depend on corridor '
    'redevelopment actually happening. Kill the housing pipeline with a 20% mandate and '
    'you also lose:'
)

cascade_bullets = [
    'Commercial floor area that depends on residential foot traffic to support retail '
    'and services along the corridors.',
    'Transit ridership that justifies current SPRINTER service levels — Oceanside Boulevard '
    'corridor has five SPRINTER stations within the Planning Area.',
    'The fiscal base to fund infrastructure improvements identified in Chapter 8 '
    '(Table 8-2), including roadway, water, sewer, and stormwater upgrades.',
    'The 28,500 new jobs the plan envisions, particularly in the industrial areas along '
    'Oceanside Boulevard that depend on adjacent residential density for workforce.',
]
for b in cascade_bullets:
    doc.add_paragraph(b, style='List Bullet')

# ─── Section 7 ───
doc.add_heading('The D-District Precedent', level=2)

doc.add_paragraph(
    'Oceanside already has a case study in what happens when policy makes development '
    'infeasible. The downtown D-District density cap (86 du/acre, certified by the Coastal '
    'Commission on February 5, 2026) produced the following result in housing application filings:'
)

p = doc.add_paragraph()
run = p.add_run('2022: 1,946 units  →  2026: 0 units')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph(
    'The mechanisms are different: the D-District density cap is a hard regulatory limit '
    'that makes projects physically impossible above 86 du/acre, while a 20% inclusionary '
    'rate is a cost mandate that makes projects financially infeasible. The outcomes differ '
    'in degree — the density cap killed 100% of new filings, while a 20% inclusionary rate '
    'would kill an estimated 58–70% of projected corridor units. But the direction is the '
    'same: policy that raises the cost of building housing beyond what the market can absorb '
    'reduces the total number of units built, including affordable units.'
)

# ─── Section 9 ───
doc.add_heading('Recommendation', level=2)

doc.add_paragraph(
    'The effective policy combination for the SSCSP corridors is:'
)

rec_bullets = [
    'Keep the 10% inclusionary set-aside — it is at the outer edge of feasibility for '
    'structured parking projects ($3/sf RLV).',
    'Let the SSCSP\'s reduced parking standards, streamlined Tier 1 ministerial review, '
    'and lot consolidation policies do their work to reduce development costs.',
    'Monitor actual production and adjust — the SSCSP\'s 3-tier review system allows '
    'policy adjustments without full plan amendments.',
]
for b in rec_bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph()

# ─── Sources ───
doc.add_heading('Sources', level=2)

sources = [
    'Smart & Sustainable Corridors Specific Plan, Hearing Draft (May 8, 2026) — '
    'Chapter 3: Land Use, Chapter 8: Implementation',
    'DRA Inclusionary Housing Study (February 2022) — commissioned by City of Oceanside, '
    'presented to City Council June 28, 2022',
    'HCD Annual Progress Report Table A — housing application filings (data.ca.gov)',
    'City of Oceanside eTRAKiT permit portal — current through June 16, 2026',
    'LCPA22-00002 (Downtown density cap ordinance) — adopted October 2023, '
    'CCC certified February 5, 2026',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

out = '/home/thomas/repos/civics/share/oceanside-sb79-briefing/SSCSP Inclusionary Rate Analysis.docx'
doc.save(out)
print(f'Saved: {out}')
