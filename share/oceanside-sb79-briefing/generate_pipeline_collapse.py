#!/usr/bin/env python3
"""Generate Word document: Oceanside 2026 Housing Pipeline Collapse."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    hs.font.name = 'Calibri'

# ─── Title ───
title = doc.add_heading('Oceanside Housing Pipeline Collapse\nJanuary–June 2026', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Data Briefing — June 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.font.italic = True

doc.add_paragraph()

# ─── Section 1: The Number ───
doc.add_heading('One Housing Application in Six Months', level=2)

p = doc.add_paragraph()
run = p.add_run(
    'Between January 1 and June 19, 2026, the City of Oceanside received one '
    'discretionary housing application.'
)
run.bold = True

doc.add_paragraph(
    'That application — DB26-00001, a 142-unit multifamily project at 1640 Oceanside '
    'Boulevard — was filed as a Density Bonus Application under state law. The project '
    'includes a minimum of 21 affordable units (15% inclusionary requirement per '
    'Oceanside\'s December 2023 ordinance). Without this state-law filing, the count '
    'would be zero.'
)

doc.add_paragraph(
    'This is a city of 183,000 people with a 6th cycle RHNA obligation of approximately '
    '7,100 units by 2029.'
)

# ─── Section 2: Historical Context ───
doc.add_heading('Historical Comparison: Discretionary Housing Filings', level=2)

doc.add_paragraph(
    'The following table shows all discretionary development applications (D-, DB-, and '
    'RD-prefix projects in the city\'s eTRAKiT permit system) filed per year, alongside '
    'total housing units reported to HCD via the Annual Progress Report.'
)

table1 = doc.add_table(rows=10, cols=4, style='Light Shading Accent 1')
table1.alignment = WD_TABLE_ALIGNMENT.CENTER
headers1 = ['Year', 'Discretionary Apps\n(D/DB/RD)', 'Total Units Filed\n(HCD APR)', 'Large Projects\n(50+ units)']
data1 = [
    ['2018', '22', '576', '1'],
    ['2019', '24', '1,090', '3'],
    ['2020', '23', '399', '1'],
    ['2021', '24', '1,698', '8'],
    ['2022', '36', '3,211', '13'],
    ['2023', '35', '1,977', '7'],
    ['2024', '34', '2,023', '10'],
    ['2025', '25 (18 housing)', '1,954', '7'],
    ['2026 (Jan–Jun)', '5 (1 housing)', '—', '0'],
]
for i, h in enumerate(headers1):
    cell = table1.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for ri, row_data in enumerate(data1):
    for ci, val in enumerate(row_data):
        cell = table1.rows[ri + 1].cells[ci]
        cell.text = val
        if ri == 8:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph()

doc.add_paragraph(
    'In 2025, the city received 18 discretionary housing applications. In the first half '
    'of 2026, it received one. The remaining 4 discretionary applications in 2026 were a '
    'police training facility, a medical office, and two Walmart expansions.'
)

# ─── Section 3: Full Housing Filings History ───
doc.add_heading('Citywide Housing Filings: Nine-Year Trend (2018–2026)', level=2)

doc.add_paragraph(
    'The following table shows all housing units filed with the city, by year and type. '
    '2018–2025 data is from HCD Annual Progress Report Table A (the same source as the '
    'D-District chart above). 2026 is reconstructed from eTRAKiT, excluding units already '
    'counted in prior years\' filings (Olive Park, North River Farms tract lots).'
)

table2 = doc.add_table(rows=11, cols=6, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
headers2 = ['Year', 'ADU', 'SFD/\nSFA', '2–4\nUnit', '5+\nUnit', 'Total']
data2 = [
    ['2018', '34', '347', '100', '78', '559'],
    ['2019', '31', '227', '22', '810', '1,090'],
    ['2020', '42', '52', '22', '283', '399'],
    ['2021', '146', '245', '57', '1,250', '1,698'],
    ['2022', '187', '200', '36', '2,774', '3,197'],
    ['2023', '149', '36', '12', '1,765', '1,962'],
    ['2024', '197', '292', '17', '1,515', '2,021'],
    ['2025', '147', '409', '9', '1,389', '1,954'],
    ['2026 (Jan–Jun)', '67', '4', '6', '142', '219'],
]
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for ri, row_data in enumerate(data2):
    for ci, val in enumerate(row_data):
        cell = table2.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
        if ri == 8:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
# Totals row
totals = ['TOTAL', '1,000', '1,812', '281', '10,006', '13,099']
for ci, val in enumerate(totals):
    cell = table2.rows[10].cells[ci]
    cell.text = val
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)

doc.add_paragraph()

doc.add_paragraph(
    'Key patterns:'
)

patterns = [
    '5+ unit projects dominate total production — averaging 1,358 units/year from '
    '2021–2025. In the first half of 2026, only 142 units were filed in this category, '
    'and that single project (DB26-00001) used AB 2011 to bypass city review.',
    'The 2022 peak (2,774 units in 5+ projects) followed the November 2021 removal of '
    'the downtown 43 du/acre cap. The October 2023 reimposition of a cap at 86 du/acre '
    'preceded the decline from 2,774 to 1,765 to 1,515 to 1,389.',
    'ADUs remained steady at ~150/year — the one type unaffected by local discretionary '
    'barriers.',
    'SFD/SFA spiked in 2025 (409 units) due to North River Farms tract lots — a '
    'pre-entitled subdivision, not new development capacity.',
    '2026 annualized (~478 units) would be the lowest production year since the city '
    'began reporting to HCD, below even the 2020 COVID year (399).',
]
for pat in patterns:
    doc.add_paragraph(pat, style='List Bullet')

doc.add_paragraph()

doc.add_paragraph(
    'Notes on 2026 data:'
)

notes = [
    '2026 HCD APR data will not be available until mid-2027. The 2026 figures are '
    'reconstructed from eTRAKiT permit and project searches.',
    'Olive Park Affordable (199 units) and North River Farms SFDs (57 lots) were excluded '
    'from 2026 because they were already counted as filings in prior years\' HCD data '
    '(DB24-00001 and the NRF subdivision, respectively).',
    'The only 5+ unit filing in 2026 is DB26-00001 (142 units at 1640 Oceanside Blvd), '
    'a Density Bonus Application with a minimum of 21 affordable units (15% inclusionary). '
    'Without this state-law filing, the 5+ unit count would be zero.',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

# ─── Section 4: The Drop ───
doc.add_heading('Total Housing Production: The Collapse in Context', level=2)

doc.add_paragraph(
    'The nine-year filing record reveals the full trajectory:'
)

table3 = doc.add_table(rows=3, cols=5, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
headers3 = ['', '2021–2022\n(avg)', '2023–2025\n(avg)', '2026*\n(annualized)', 'Change']
data3 = [
    ['Total units filed', '2,448', '1,979', '~478', '-80%'],
    ['5+ unit projects', '2,012', '1,556', '~310', '-85%'],
]
for i, h in enumerate(headers3):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
for ri, row_data in enumerate(data3):
    for ci, val in enumerate(row_data):
        cell = table3.rows[ri + 1].cells[ci]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
        if ci == 4:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)

doc.add_paragraph(
    '* 2026 is Jan–Jun, annualized. 2026 data from eTRAKiT, deduped against prior-year '
    'HCD filings.',
    style='List Bullet'
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    'At 219 units filed in 5.5 months (annualized ~478), 2026 would be the lowest '
    'filing year since Oceanside began reporting to HCD — below even the 2020 COVID '
    'year (399 units). And the only 5+ unit project filed used state law to bypass '
    'local review entirely.'
)
run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run(
    'The RHNA target requires roughly 890 new units per year. At the current filing '
    'pace, Oceanside will produce less than half its annual obligation, accumulating '
    'a deficit that makes the 2029 deadline unreachable.'
)
run.bold = True

# ─── Section 5: The Lag Effect ───
doc.add_heading('The Lag Effect: Building Permits Will Get Worse', level=2)

doc.add_paragraph(
    'Building permits lag discretionary approvals by 1–2 years. The multi-family and '
    'mid/high-rise building permits issued in 2025 came from projects approved in 2022–2024:'
)

lag_examples = [
    'Modera Neptune (BLD HIGH RISE, BLDG25-0374): 360 units, approved 2022 as RD22-00005.',
    'Coast Villas (BLD MID RISE, BLDG25-1204): 56 units, approved 2024 as DB25-00002.',
    'Pacifica townhomes (30 BLD MULTI FAMILY permits): 164 units across 12 phases, '
    'approved 2022 as D22-00013.',
]
for ex in lag_examples:
    doc.add_paragraph(ex, style='List Bullet')

doc.add_paragraph()

doc.add_paragraph(
    'With only one discretionary housing application in 2026, there will be few or no '
    'multi-family building permits to pull in 2027–2028. The 2025 building permit numbers '
    'represent the last wave of construction from the pre-collapse pipeline. When those '
    'projects finish, the building permit pipeline will mirror the discretionary freeze.'
)

# ─── Section 6: Why ───
doc.add_heading('What Caused the Collapse', level=2)

doc.add_paragraph(
    'Three policy actions converged to freeze the discretionary pipeline:'
)

causes = [
    ('D-District Density Cap (CCC Certified February 5, 2026). ',
     'The downtown density cap of 86 du/acre, adopted by the City Council in October 2023 '
     'and certified by the Coastal Commission on February 5, 2026, eliminated the city\'s '
     'most productive housing area. Downtown D-District filings went from 1,946 units in '
     '2022 to zero in 2026. No new D-District housing project has been filed since '
     'certification.'),
    ('SB 79 Exemption Ordinance (Adopted June 3, 2026). ',
     'The city adopted an ordinance deferring SB 79 transit-oriented housing compliance '
     'to 2032, using fabricated "walking path" findings to reclassify parcels near transit '
     'as non-qualifying. This froze development potential on remaining non-downtown parcels '
     'that could have absorbed displaced demand.'),
    ('SSCSP Not Yet Adopted. ',
     'The Smart & Sustainable Corridors Specific Plan, which would rezone 1,172 acres '
     'of corridor land to allow 8,300 new housing units, has not yet been adopted. The '
     'hearing is scheduled for June 24, 2026. Until adoption, corridor parcels remain '
     'zoned for low-density residential and light industrial uses (3.9–7.25 du/acre), '
     'making multifamily development physically impossible under existing zoning.'),
]

for title_text, body_text in causes:
    p = doc.add_paragraph()
    run = p.add_run(title_text)
    run.bold = True
    p.add_run(body_text)

# ─── Section 7: The Only Housing Filing Used State Law ───
doc.add_heading('The Only Housing Filing Bypassed City Process', level=2)

doc.add_paragraph(
    'DB26-00001 (142 units at 1640 Oceanside Boulevard) was filed on June 16, 2026 — '
    'eight days before the SSCSP adoption hearing. The project is located on a parcel '
    'currently zoned RE-B (Residential Estate Bonus, 3.9 du/acre maximum). At 142 units '
    'on approximately 1.4 acres, the project proposes roughly 102 du/acre — 26 times '
    'the current zoning allows.'
)

doc.add_paragraph(
    'The project is filed as a Density Bonus Application (eTRAKiT type "DB"), invoking '
    'state density bonus law (Government Code §65915). Under Oceanside\'s inclusionary '
    'housing ordinance (Chapter 14C, updated December 2023 to 15%), the project must '
    'include a minimum of 21 deed-restricted affordable units at lower income levels '
    '(≤80% AMI). The actual affordable count may be higher if the developer offered '
    'additional affordable units in exchange for density bonus concessions.'
)

p = doc.add_paragraph()
run = p.add_run(
    'The fact that the only housing project filed in Oceanside in 2026 had to use state '
    'law to avoid local review is itself evidence that the local regulatory environment '
    'has become hostile to housing production.'
)
run.bold = True

# ─── Section 8: Data Verification ───
doc.add_heading('Data Verification', level=2)

doc.add_paragraph(
    'The data in this briefing was cross-verified across multiple sources:'
)

verification = [
    'Discretionary projects: City of Oceanside eTRAKiT permit portal, project search '
    '(D-, DB-, RD-prefix applications). All projects for 2018–2026 were individually '
    'verified by direct URL query.',
    'Citywide housing filings (2018–2025): HCD Annual Progress Report Table A, '
    'APP_SUBMIT_DT field — same source and methodology as the D-District analysis.',
    '2026 citywide filings: Reconstructed from eTRAKiT discretionary project search '
    '(D/DB/RD-prefix) and building permit search (BLDG-prefix, 20,702 permits scraped '
    '2020–2026). Deduped against prior-year HCD filings to avoid double-counting '
    'Olive Park (DB24-00001) and North River Farms tract lots.',
    'HCD Annual Progress Report: Table A data from data.ca.gov, filtered to '
    'JURIS_NAME=OCEANSIDE, covering 2018–2025. 2026 APR data will not be available '
    'until mid-2027.',
    'Cross-reference: 45 of 50 large APR projects (50+ units, 2018–2025) were matched '
    'to eTRAKiT discretionary applications by Assessor Parcel Number (APN), confirming '
    'both data sources track the same projects.',
    'The 5 unmatched projects were traced to APN discrepancies between APR and eTRAKiT '
    '(different parcel numbers for the same site), not missing projects.',
]
for v in verification:
    doc.add_paragraph(v, style='List Bullet')

# ─── Section 9: Implications ───
doc.add_heading('Implications', level=2)

implications = [
    'RHNA compliance: At the current filing rate (~478 units/year annualized, of which '
    'only 21 are affordable), Oceanside will produce barely half its annual RHNA obligation '
    'of ~890 units — accumulating a deficit that makes the 2029 deadline unreachable.',
    'Builder\'s Remedy exposure: If the city fails to maintain adequate housing element '
    'progress, it risks losing its certified housing element status, triggering Builder\'s '
    'Remedy (Government Code 65589.5(d)), which allows developers to bypass local zoning '
    'entirely.',
    'SSCSP urgency: The corridor plan is now the only mechanism that could restore '
    'discretionary housing capacity. Delaying adoption or imposing additional cost '
    'mandates (such as a 20% inclusionary rate) would extend the production freeze.',
    'The lag effect: Even if the SSCSP is adopted on June 24 and developers begin filing '
    'immediately, those projects will not pull building permits until 2027–2028. The '
    'construction gap is already locked in for at least 18 months.',
    'State enforcement: HCD has authority to revoke housing element certification for '
    'jurisdictions that adopt policies inconsistent with their housing element. The '
    'combination of the D-District cap, SB 79 deferment, and pipeline collapse may '
    'constitute grounds for decertification.',
]
for imp in implications:
    doc.add_paragraph(imp, style='List Bullet')

doc.add_paragraph()

# ─── Sources ───
doc.add_heading('Sources', level=2)

sources = [
    'City of Oceanside eTRAKiT permit portal — project search and permit search, '
    'individually verified through June 19, 2026',
    'HCD Annual Progress Report Table A (data.ca.gov) — 2018–2025 filing data',
    'LCPA22-00002 (Downtown density cap ordinance) — adopted October 2023, '
    'CCC certified February 5, 2026',
    'Oceanside SB 79 exemption ordinance — adopted June 3, 2026',
    'Smart & Sustainable Corridors Specific Plan, Hearing Draft (May 8, 2026)',
    'DB26-00001 application record — eTRAKiT, filed June 16, 2026',
    'AB 2011 (Affordable Housing and High Road Jobs Act of 2022) — '
    'Government Code 65912.100 et seq.',
]
for s in sources:
    doc.add_paragraph(s, style='List Bullet')

out = '/home/thomas/repos/civics/share/oceanside-sb79-briefing/Oceanside 2026 Housing Pipeline Collapse.docx'
doc.save(out)
print(f'Saved: {out}')
