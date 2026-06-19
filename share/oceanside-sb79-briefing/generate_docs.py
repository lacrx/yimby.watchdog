#!/usr/bin/env python3
"""Generate Word documents for Oceanside SB 79 / density cap activist briefing."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))


def style_doc(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)


def add_title(doc, text, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(11)
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        run2.italic = True


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    return h


def add_bold_para(doc, bold_text, normal_text=""):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def add_callout(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.right_indent = Cm(1)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(0x8B, 0x00, 0x00)
    return p


# ============================================================
# DOCUMENT 1: TOP-LINE SUMMARY (1-2 pages)
# ============================================================
def build_summary():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Oceanside's War on Housing",
              "Top-Line Summary for Activists | June 2026")

    doc.add_paragraph()
    add_heading(doc, "What's Happening", level=1)
    doc.add_paragraph(
        "Oceanside built a legal machine to kill housing across the entire city. Not just "
        "downtown. All seven stations. It runs on three tracks:"
    )
    add_bullet(doc, "86 du/acre downtown, slashing allowable housing by ~40%.",
               "Density cap: ")
    add_bullet(doc, "A Coastal Commission ruling (Feb 2026) that lets the city kill density "
               "bonus projects by citing 'viewsheds' and 'community character.' Whatever works.",
               "Coastal Act shield: ")
    add_bullet(doc, "A citywide evasion ordinance (June 3, 2026 council hearing, Planning Commission 7-0) designed "
               "to gut SB 79 at every transit stop. SANDAG parcel data shows ~18,800 parcels within half a mile "
               "of the city's 8 stations. The walking path exemption — the most documented tool — covers only 710 "
               "of those (3.8%), concentrated around Crouch Street and Civic Center. OTC, the city's biggest "
               "station at 130+ trains/day, has 4,572 parcels in its SB 79 zone and just 2 in the walking path "
               "shapefile. The city claims deferments, rent control misclassification, habitat exclusions, and "
               "a planned Alternative Plan cover the rest — but no public data confirms this. The gap between "
               "what's documented and what's claimed is the story.",
               "SB 79 sabotage: ")

    add_heading(doc, "This Is Citywide", level=2)
    doc.add_paragraph(
        "SB 79 covers all 7 transit stops. The evasion ordinance targets them all:"
    )
    add_bullet(doc, "~130 trains/day. Tier 1 threshold is 72. "
               "The city classified it Tier 2 anyway, dropping required density from 100–120 to "
               "80–100 du/acre. 130 is bigger than 72. This is not a judgment call.",
               "OTC: ")
    add_bullet(doc, "Coast Highway, Crouch Street, El Camino Real, Rancho Del Oro, and the rest "
               "(68 trains/day, Tier 2 threshold 48). The city fabricated walking path definitions, "
               "invented habitat exclusions that don't exist in state law, and reclassified Good Cause "
               "eviction as 'rent control' to carve out parcels around every station.",
               "Sprinter stops: ")
    doc.add_paragraph(
        "Combined SB 79 zones around all 7 stops could hold 13,000–32,000 units. That's enough "
        "to satisfy Oceanside's remaining RHNA 2–10 times over. The evasion ordinance exists to "
        "zero that out."
    )

    add_heading(doc, "Downtown Is Already Dead", level=2)
    doc.add_paragraph(
        "Since the Coastal Commission certified the density cap in February 2026, zero new "
        "housing projects have been approved downtown. Zero. The Rodeway Inn, 290 units, "
        "biggest project in the pipeline, is being pulled. 801 Mission (230 units) sat frozen "
        "for months before scraping through in late May."
    )

    add_heading(doc, "The Stakes", level=1)

    add_heading(doc, "Hard costs: what the city will actually spend", level=2)
    add_table(doc,
              ["Category", "Estimate", "Basis"],
              [
                  ["Outside counsel + litigation defense", "$2M–$5M",
                   "SB 79 suit, HAA claims, density bonus cases. Based on Huntington Beach, La Cañada Flintridge comparables"],
                  ["Court-ordered attorney's fees (mandatory on loss)", "$500K–$1.5M",
                   "Losing city pays plaintiff's lawyers. CalHDF collected $1.26M from La Cañada Flintridge alone"],
                  ["AG fines ($10K–$50K/month)", "$240K–$1.8M",
                   "Accumulates monthly once AG acts. Huntington Beach at $50K/month as of May 2026"],
                  ["HAA penalties on denied projects", "$830K–$4.15M",
                   "83-unit denial: $10K/unit minimum. 5x bad-faith multiplier possible"],
                  ["Total hard costs", "$4M–$16M",
                   "Checks from the General Fund. Same pot that pays for police, fire, parks, libraries"],
              ])

    add_heading(doc, "Opportunity costs: revenue that never arrives", level=2)
    doc.add_paragraph(
        "These are not bills in the mail. They're fees and taxes the city would collect if housing "
        "gets built. They won't collect them if the pipeline stays frozen. We include them because "
        "they're real consequences, but we label them for what they are: projections based on what "
        "comparable projects generated elsewhere. Not guaranteed losses."
    )
    add_table(doc,
              ["Category", "Estimate", "Basis"],
              [
                  ["Building permit fees below trend", "$6M–$12M",
                   "3–5 years of declining fee revenue as pipeline dries up"],
                  ["Developer impact fees not collected", "$10M–$40M",
                   "500–1,000 chilled units that won't pay $20K–$40K/unit in fees"],
                  ["Property tax revenue delayed", "$17M–$32M",
                   "700+ units that would generate $8K–$15K/yr in property tax, delayed 3+ years"],
                  ["Total opportunity cost", "$28M–$70M",
                   "Revenue the city misses out on. Real but speculative, depends on what would have been built"],
              ])

    doc.add_paragraph(
        "Combined: $32M–$86M. But the number you can hold the council to is $4M–$16M. "
        "That's taxpayer cash going to lawyers instead of services."
    )

    add_heading(doc, "The Endgame", level=1)
    doc.add_paragraph(
        "SB 79 takes effect July 1, 2026. By-right housing at 100–120 du/acre within half a "
        "mile of major transit. 80–100 near lighter rail. All 7 Oceanside stations qualify. The "
        "OTC alone runs ~130 trains a day against a Tier 1 threshold of 72. 'By-right' means the "
        "city doesn't get a vote. No coastal impact analysis. No viewshed denial. No council "
        "hearing. Meet the objective standards, get the permit."
    )
    doc.add_paragraph(
        "The city is betting it can stall SB 79 citywide long enough to scare developers off "
        "permanently. CalHDF, YIMBY Law, and Californians for Homeownership, backed by a "
        "six-org coalition, are getting ready to sue."
    )
    doc.add_paragraph(
        "Most likely outcome: the city fights 2–3 years, spends tens of millions, and loses. "
        "No California city has ever won this fight. Not one."
    )

    add_callout(doc, "Bottom line: $4M–$16M in real legal costs from the General Fund. "
                "$28M–$70M more in lost fees and taxes on top. Every city that's tried this has "
                "lost. Only question is how much burns first.")

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run("See companion documents for full legal strategy details and cost breakdown.")
    run.italic = True

    doc.save(os.path.join(OUT_DIR, "1 - Top-Line Summary.docx"))
    print("Built: 1 - Top-Line Summary.docx")


# ============================================================
# DOCUMENT 2: THE LEGAL STRATEGY (detailed)
# ============================================================
def build_legal_strategy():
    doc = Document()
    style_doc(doc)
    add_title(doc, "Oceanside's Eight-Layer Legal Strategy",
              "How the City Built a Housing Blockade | June 2026")

    add_heading(doc, "Overview", level=1)
    doc.add_paragraph(
        "Oceanside hasn't openly defied state housing law like Huntington Beach. Instead, it "
        "has built the most sophisticated housing obstruction system in California. Eight "
        "mutually reinforcing mechanisms, each deniable on its own, lethal in combination. "
        "The city never says \"no.\" It says \"not this way.\""
    )

    add_heading(doc, "The Eight Layers", level=1)

    layers = [
        ("Layer 1: Downtown Density Cap (86 du/acre)",
         "Reduced from unlimited (2019–2023) to 86 du/acre. At Oceanside land prices, most "
         "projects need 120–200 du/acre to pencil with any affordable component. The cap alone "
         "eliminates ~40% of downtown's housing capacity. About 600–980 units that would "
         "otherwise be built.",
         "SB 330 downzoning challenge; baseline question (was it 43 or unlimited?)"),

        ("Layer 2: Coastal Act Shield",
         "The Coastal Commission certified a modification (Feb 5, 2026) codifying that density "
         "bonus law \"shall not supersede\" the Coastal Act. The city can now deny density bonus "
         "waivers (height, setbacks, parking) citing viewsheds, public access, habitat, or "
         "\"community character.\" No objective standards exist for what counts as a coastal impact. "
         "This transforms every density bonus project from a by-right approval into a discretionary "
         "gamble.",
         "Untested against SB 79's by-right mechanism; novel legal question"),

        ("Layer 3: SB 79 Tier Misclassification",
         "The Oceanside Transit Center has ~130 commuter rail trains/day. SB 79's Tier 1 threshold "
         "is 72 trains. The city classified it as Tier 2 (48–71 trains), reducing required density "
         "from 100–120 to 80–100 du/acre. This is not an interpretation. 130 > 72 is arithmetic. "
         "CalHDF called it one of the most brazen fabrications in the state, comparable to Woodside "
         "declaring itself a mountain lion habitat.",
         "Indefensible. 90–95% probability city loses. SANDAG's official TOD map (forthcoming) "
         "will likely designate OTC as Tier 1."),

        ("Layer 4: Walking Path Fabrication",
         "SB 79 lets cities exempt sites without a \"walking path\" within one mile of a transit "
         "stop. Oceanside defined \"walking path\" as requiring continuous paved sidewalk with no "
         "gaps. Then ran the analysis. 710 parcels tested. Zero passed. Not one. The city's own "
         "sidewalk inventory (produced under PRA) shows 71% of road segments near stations have "
         "sidewalks. 1,229 out of 1,726 segments. But any gap in the chain breaks the \"path.\" "
         "Critical context: SANDAG parcel data shows ~18,800 total parcels within half a mile of "
         "the city's 8 stations. The walking path shapefile covers 710 — just 3.8%. It concentrates "
         "around Crouch Street (~511 parcels) and Civic Center (~89). OTC has 4,572 parcels in its "
         "SB 79 zone and only 2 in the walking path dataset. The city claims deferments and other "
         "tools cover the remaining 96%, but no parcel-level data confirming this has been produced.",
         "SB 79 says \"walking path,\" not \"continuous paved sidewalk.\" Planning staff admitted "
         "at the May 18 hearing: \"there is no definition of walking path in the statute, so the "
         "city created our own definition.\" Vehicle Code 21954 permits pedestrians on roadways. "
         "AB 2147 author: \"People should not be penalized for decades of infrastructure neglect.\" "
         "CalHDF identified specific parcels that abut stations and were still exempted. The gap "
         "between documented exemptions (3.8%) and claimed coverage (100%) is itself evidence: "
         "either the other tools have massive holes, or the city has parcel data it hasn't disclosed."),

        ("Layer 5: Habitat/Wetland Exclusions",
         "General Plan Policy 5-12 (habitat preserves, wetland buffers) was used to exclude sites "
         "from SB 79 coverage. SB 79 contains no such exception. State law preempts local policy "
         "under the California Constitution (Art. XI, § 7).",
         "No legal basis in SB 79. Clear state preemption."),

        ("Layer 6: Good Cause ≠ Rent Control",
         "The city classified Good Cause eviction protections (AB 1482, Civil Code Ch. 2) as "
         "\"rent control\" (Civil Code Ch. 2.7) to exempt additional properties from SB 79. These "
         "are legally distinct categories under Costa-Hawkins. The California Supreme Court held "
         "in Briggs v. Eden Council that different statutory terms are presumed to mean different "
         "things.",
         "Clear legal error. Costa-Hawkins distinction is well-settled."),

        ("Layer 7: De Novo Appeal Ordinance",
         "Adopted June 2025 (4-1). Allows the City Council to conduct full de novo review of any "
         "appealed Planning Commission approval. SB 330 limits projects to 5 hearings total. Each "
         "use of de novo appeal on a compliant project is independently actionable under the Housing "
         "Accountability Act.",
         "SB 330 § 65905.5(a) five-hearing cap. Each use = separate HAA violation."),
    ]

    for title, desc, vuln in layers:
        add_heading(doc, title, level=2)
        doc.add_paragraph(desc)
        add_bold_para(doc, "Legal vulnerability: ", vuln)

    add_heading(doc, "Layer 8: NCTD Board Coordination", level=2)
    doc.add_paragraph(
        "Deputy Mayor Eric Joyce sits on both the Oceanside City Council and the NCTD Board. "
        "At NCTD in April 2026, he authored a motion (passed unanimously) to oppose any "
        "legislation expanding SB 79 and delay its effective dates. The transit board member "
        "who could advocate for service levels that strengthen Tier 1 classification is instead "
        "lobbying to kill the law entirely. This dual-role coordination between local government "
        "and the regional transit agency is a distinct obstruction pattern not visible in city "
        "records alone."
    )

    add_heading(doc, "How the Layers Work Together", level=1)
    doc.add_paragraph(
        "Each layer provides fallback if others fail. If the density cap is challenged under "
        "SB 330, the Coastal Act shield blocks density bonus overrides. If SB 79 overrides "
        "the Coastal Act, the Tier misclassification reduces required density. If the Tier "
        "classification is corrected, the walking path fabrication excludes parcels. If all "
        "else fails, the de novo appeal ordinance lets the council add unlimited hearings to "
        "any project that gets through."
    )
    doc.add_paragraph(
        "This is defensive depth rare among California cities. Most resisting cities have one "
        "or two blocking tools. Oceanside has eight, including cross-jurisdictional coordination "
        "through NCTD."
    )

    add_heading(doc, "The Procedural Defect: Independently Fatal", level=1)
    doc.add_paragraph(
        "Beyond the substantive violations, the SB 79 evasion ordinance has a procedural defect "
        "that may be independently invalidating: the City Council hearing was noticed before the "
        "Planning Commission voted on May 18, 2026. The PC recommendation was not posted as of "
        "May 20. The vote field was left blank. This violates Government Code §§ 65854 et seq. "
        "and Environmental Defense Project v. County of Sierra (158 Cal.App.4th 877). The "
        "ordinance may be voidable on procedural grounds alone, regardless of whether the "
        "substantive violations are proven."
    )

    doc.save(os.path.join(OUT_DIR, "2 - The Legal Strategy.docx"))
    print("Built: 2 - The Legal Strategy.docx")


# ============================================================
# DOCUMENT 3: COSTS AND CONSEQUENCES
# ============================================================
def build_costs():
    doc = Document()
    style_doc(doc)
    add_title(doc, "What This Will Cost Oceanside",
              "Financial and Human Consequences | June 2026")

    add_heading(doc, "About Honesty", level=1)
    doc.add_paragraph(
        "This document separates hard costs from opportunity costs. Hard costs are checks the "
        "city writes. Legal fees. Court-ordered penalties. Attorney's fees awards. Opportunity "
        "costs are revenue the city misses because projects never get built. Both are real. "
        "They're different kinds of real. Mixing them into one headline number is the same "
        "dishonesty we're calling the city out for, and we're not doing it."
    )

    add_heading(doc, "Hard Costs: What the City Will Actually Spend", level=1)
    doc.add_paragraph(
        "These come straight from the General Fund. Same pot that pays for police, fire, parks, "
        "libraries. Every dollar here is a dollar not spent on services."
    )
    add_table(doc,
              ["Front", "Low Estimate", "High Estimate", "Source / Basis"],
              [
                  ["SB 79 litigation defense (outside counsel)", "$800K", "$2M",
                   "La Cañada Flintridge spent ~$2M on simpler case"],
                  ["Plaintiff attorney's fees (mandatory, losing city pays)", "$500K", "$1.5M",
                   "CalHDF collected $1.26M from La Cañada Flintridge"],
                  ["AG fines ($10K–$50K/month)", "$240K", "$1.8M",
                   "Huntington Beach now paying $50K/month; accumulates for 12–36 months"],
                  ["83-unit EIR denial, HAA penalties ($10K/unit)", "$830K", "$4.15M",
                   "Statutory minimum; 5x bad-faith multiplier if pattern established"],
                  ["Ongoing outside counsel / closed session advisory", "$200K/yr", "$500K/yr",
                   "Standard rate for municipal housing litigation; city is already incurring this"],
                  ["Total hard costs", "$4M", "$16M",
                   "Checks written from the General Fund"],
              ])
    doc.add_paragraph(
        "Low-end estimates are conservative. High-end includes bad-faith multipliers. The "
        "density bonus waiver denial exposure ($7M–$86M per the full legal forecast) isn't in "
        "this table because no waivers have been denied yet. That exposure exists but hasn't "
        "triggered. If it does, this table gets much worse."
    )

    add_heading(doc, "What Gets Cut", level=2)
    doc.add_paragraph(
        "Oceanside runs a slim General Fund surplus. About $9.9M a year. CalPERS pension "
        "obligations are eating into it at ~$2.5M/year and climbing. $4M in legal costs wipes "
        "out nearly half that surplus. At $16M, the city borrows against future budgets."
    )
    doc.add_paragraph(
        "What goes first: parks, arts, libraries, community programs, infrastructure "
        "maintenance. Discretionary spending always dies first."
    )

    add_heading(doc, "Opportunity Costs: Revenue That Never Arrives", level=1)
    doc.add_paragraph(
        "These are not bills in the mail. They're fees and taxes the city would collect if "
        "housing gets built. They won't collect them if the pipeline stays frozen. We include "
        "them because they're real consequences, but we label them for what they are: projections "
        "based on what comparable projects generated elsewhere. Not guaranteed losses."
    )
    add_table(doc,
              ["Category", "Low", "High", "How estimated"],
              [
                  ["Building permit fees below trend", "$6M", "$12M",
                   "3–5 years of declining fee revenue as pipeline dries up"],
                  ["Developer impact fees not collected", "$10M", "$40M",
                   "500–1,000 chilled units × $20K–$40K/unit (based on city's own fee schedule)"],
                  ["Property tax revenue delayed", "$17M", "$32M",
                   "700+ units × $8K–$15K/yr × 3 years (county assessor rate on comparable projects)"],
                  ["Affordable housing in-lieu fees", "$3M", "$8M",
                   "Projects that can't pencil at 86 du/acre without density bonus"],
                  ["Total opportunity cost", "$28M", "$70M",
                   "Speculative but grounded in city fee schedules and comparable project data"],
              ])

    add_heading(doc, "Combined Picture", level=1)
    add_table(doc,
              ["Category", "Low", "High", "Type"],
              [
                  ["Hard costs (legal + penalties)", "$4M", "$16M", "Money the city will spend"],
                  ["Opportunity costs (lost revenue)", "$28M", "$70M", "Revenue the city won't collect"],
                  ["Combined", "$32M", "$86M", "Hard + opportunity"],
              ])
    doc.add_paragraph(
        "Talk to the council or press? Lead with hard costs. $4M–$16M is unambiguous taxpayer "
        "money going to lawyers. Mention the $28M–$70M in opportunity costs separately. Label "
        "them as projections. More credible than one inflated headline number. Harder to dismiss, too."
    )

    add_heading(doc, "What Compliance Would Cost", level=1)
    add_table(doc,
              ["Action", "Cost"],
              [
                  ["Correct SB 79 tier to Tier 1", "$0 (staff time)"],
                  ["Accept density bonus waivers", "$0"],
                  ["Drop SB 79 evasion ordinance", "$0"],
                  ["Process compliant projects on time", "$0"],
                  ["Total cost of compliance", "~$0"],
              ])

    add_callout(doc, "The hard-cost difference between fighting and complying is $4M–$16M. "
                "Including lost revenue, $32M–$86M. The cost of just following the law is zero.")

    add_heading(doc, "Who Pays for Affordable Housing", level=1)
    doc.add_paragraph(
        "The density bonus is what makes affordable units pencil in market-rate projects. "
        "Developers include affordable units and get extra height and density in return. Kill "
        "the density bonus through the Coastal Act shield and there's no reason to include "
        "affordable units beyond the base 15% inclusionary rate. That rate only kicks in at "
        "7+ units anyway."
    )
    doc.add_paragraph(
        "Oceanside is at 11% of its lower-income RHNA target. Midpoint of the cycle. The city "
        "needs ~1,764 more very-low and low-income units by 2029. Without a functioning density "
        "bonus, the only way to build affordable housing is direct public subsidy. The city "
        "isn't providing that. The $16.5M surplus from FY2023–24 went to river restoration and "
        "a park. Not housing."
    )
    add_callout(doc, "The people most hurt by this strategy are low-income renters. They need "
                "the affordable units that density bonus projects create. The city says it supports "
                "affordable housing while destroying the only mechanism that produces it.")

    add_heading(doc, "The Capitulation Timeline", level=1)
    add_table(doc,
              ["Date", "Event"],
              [
                  ["Summer 2026", "CalHDF files SB 79 suit. AG sends letter."],
                  ["Fall 2026", "Preliminary injunction hearing on Tier classification."],
                  ["Mid-2027", "Trial court ruling. City likely loses (130 > 72)."],
                  ["Late 2027", "City appeals. Must post appeal bond ($5M–$14M)."],
                  ["Mid-2028", "Most likely capitulation. Legal bills hit $5M+, AG fines 18+ months."],
                  ["2029", "RHNA cycle ends. If still fighting: possible court-appointed monitor."],
              ])

    doc.save(os.path.join(OUT_DIR, "3 - Costs and Consequences.docx"))
    print("Built: 3 - Costs and Consequences.docx")


# ============================================================
# DOCUMENT 4: THE OTC MISCLASSIFICATION
# ============================================================
def build_otc_misclassification():
    doc = Document()
    style_doc(doc)
    add_title(doc, "How Oceanside Hid 62 Trains",
              "The OTC Tier Misclassification | June 2026")

    add_heading(doc, "The Short Version", level=1)
    doc.add_paragraph(
        "Oceanside Transit Center gets at least 130 commuter rail trains every "
        "weekday. SB 79's Tier 1 threshold is 72. Oceanside classified OTC as Tier 2."
    )
    doc.add_paragraph(
        "130 is greater than 72. Staff knew. Their own presentations reveal the "
        "method: count only one operator, bury three others, and never show council "
        "the aggregate."
    )

    add_heading(doc, "The Actual Count", level=1)
    doc.add_paragraph(
        "Four operators run commuter rail through OTC. HCD's March 20, 2026 advisory "
        "to metropolitan planning organizations is unambiguous: aggregate ALL services "
        "across both directions at each station. \"This approach reflects the overall "
        "level of transit access and frequency experienced by riders at a given stop.\""
    )
    add_table(doc,
              ["Operator", "Service", "Weekday Runs", "Notes"],
              [
                  ["NCTD", "Sprinter", "68",
                   "34 each direction. Hybrid rail on fixed track."],
                  ["NCTD", "COASTER", "30",
                   "Commuter rail, San Diego to Oceanside corridor"],
                  ["Amtrak", "Pacific Surfliner", "26",
                   "San Diego to LA/San Luis Obispo. Extra Friday service."],
                  ["Metrolink", "Inland Empire-OC Line", "6",
                   "Commuter rail, terminates at Oceanside"],
                  ["", "TOTAL", "130+",
                   "Tier 1 bar: 72. Clears it by 80%."],
              ])
    doc.add_paragraph(
        "Source: YIMBY Dems and CalHDF joint SANDAG letter, May 29, 2026. "
        "Numbers confirmed against published GTFS schedules."
    )

    add_heading(doc, "What Council Actually Saw", level=1)
    doc.add_paragraph(
        "Staff briefed council twice: January 14, 2026 (Item 18) and April 8 "
        "(Item 31). Both decks carry an identical Key Definitions slide. Word for "
        "word, here is what it says about frequency:"
    )
    add_callout(doc,
        "\"Note: Sprinter currently operates with 34 trains in each direction, or 68 "
        "trains total on a typical weekday.\n\n"
        "The Sprinter line is technically neither a commuter rail line or light rail "
        "transit service. Rather, it's considered a 'hybrid' line which is undefined "
        "in the legislation.\"")
    doc.add_paragraph(
        "Full stop. No COASTER. No Surfliner. No Metrolink. No aggregation rule. "
        "Staff showed a single-operator snapshot and let the math speak for itself: "
        "68 is less than 72, so OTC looks like Tier 2."
    )

    add_heading(doc, "The Selective Lens", level=2)
    add_table(doc,
              ["What staff displayed", "What staff omitted"],
              [
                  ["Sprinter: 68 runs/day", "COASTER: 30 runs/day"],
                  ["\"Hybrid line, undefined in legislation\"",
                   "Pacific Surfliner: 26 runs/day"],
                  ["Implication: OTC sits below 72",
                   "Metrolink: 6 runs/day"],
                  ["", "Aggregate: 130+ runs/day"],
                  ["", "HCD March 2026 advisory: count ALL operators"],
              ])
    doc.add_paragraph(
        "Staff didn't miscount. They picked which runs to display. "
        "Sprinter alone (68) sits four below the threshold. Toss in COASTER and "
        "you're at 98. Toss in every operator and you're at 130. The gap between "
        "68 and 130 is 62 phantom trains that never appeared on a single council slide."
    )

    add_heading(doc, "The 'Hybrid' Smokescreen", level=1)
    doc.add_paragraph(
        "Both presentations label Sprinter \"hybrid\" and flag the term as \"undefined "
        "in the legislation.\" Implication: maybe it shouldn't count at all. Two "
        "problems with that."
    )
    doc.add_paragraph(
        "Problem one: strip Sprinter out entirely. COASTER, Surfliner, and Metrolink "
        "still deliver 62 daily runs. Below Tier 1, yes. But staff never mentioned those "
        "62 either, because naming them would expose OTC as a multi-operator hub, not a "
        "lone Sprinter stop."
    )
    doc.add_paragraph(
        "Problem two: Sprinter does qualify. SB 79 defines commuter rail as "
        "\"a public rail transit service not meeting the standards for heavy rail "
        "or light rail, excluding California High-Speed Rail and Amtrak Long Distance "
        "Service.\" Sprinter is public rail. It isn't heavy rail. NCTD itself called "
        "it light rail in board resolutions 03-11, 09-16, and 15-08. Whichever box "
        "you put it in, the station clears Tier 1."
    )

    add_heading(doc, "A Timeline of Omissions", level=1)
    add_table(doc,
              ["Date", "What happened"],
              [
                  ["Jan 14, 2026", "Staff debut SB 79 for council. One slide. Sprinter "
                   "only. 68 trains."],
                  ["Mar 20, 2026", "HCD publishes aggregation advisory. Every operator, "
                   "both directions, same station."],
                  ["Apr 8, 2026", "Second council briefing. Same deck. Same slide. Same "
                   "lone Sprinter number. Zero acknowledgment of HCD guidance."],
                  ["May 18, 2026", "Planning Commission rubber-stamps SB 79 ordinance 7-0. "
                   "OTC locked in as Tier 2."],
                  ["May 29, 2026", "Six-org coalition (YIMBY Dems, CalHDF, YIMBY Law, "
                   "CA YIMBY, Circulate, Californians for Homeownership) sends SANDAG "
                   "the real math. 130 trains. Four operators. Tier 1."],
              ])
    doc.add_paragraph(
        "Staff had two months between HCD's advisory and the Planning Commission vote. "
        "Updating a PowerPoint slide was apparently too heavy a lift."
    )

    add_heading(doc, "Density at Stake", level=1)
    doc.add_paragraph(
        "This is not a technicality. The tier controls how much gets built near OTC."
    )
    add_table(doc,
              ["", "Tier 2 (staff classification)", "Tier 1 (correct)"],
              [
                  ["Max density (within 1/4 mi)", "100 du/ac", "120 du/ac"],
                  ["Max height (within 1/4 mi)", "65 ft", "75 ft"],
                  ["Max FAR (within 1/4 mi)", "3.0", "3.5"],
                  ["Max density (1/4-1/2 mi)", "80 du/ac", "100 du/ac"],
                  ["Max height (1/4-1/2 mi)", "55 ft", "65 ft"],
                  ["Max FAR (1/4-1/2 mi)", "2.5", "3.0"],
                  ["Adjacent intensifier", "+20ft, +40 du/ac, +1.0 FAR",
                   "+20ft, +40 du/ac, +1.0 FAR"],
              ])
    doc.add_paragraph(
        "Within a quarter mile, Tier 2 caps density at 100 du/acre and height at "
        "65 feet. Tier 1 allows 120 du/acre and 75 feet. That 20% haircut is the "
        "margin on every deal. At Oceanside land prices, projects pencil with "
        "affordable units at 120. They don't at 100. Drop to 80 in the quarter-to-"
        "half-mile ring and nothing pencils."
    )

    add_heading(doc, "Litigation Exposure", level=1)
    doc.add_paragraph(
        "This layer is Oceanside's weakest. Other parts of the legal strategy "
        "demand statutory interpretation. This one demands a calculator."
    )
    add_bullet(doc, "SANDAG's forthcoming TOD map carries a rebuttable presumption "
               "of validity (Gov. Code 65912.160(f)). The six-org coalition has asked "
               "SANDAG to designate OTC as Tier 1. If SANDAG complies, Oceanside would "
               "be arguing against its own regional planning body.")
    add_bullet(doc, "CalHDF put it bluntly: the misclassification is \"one of the most "
               "brazen fabrications in the state,\" on par with Woodside's mountain-lion "
               "stunt.")
    add_bullet(doc, "Starting January 1, 2027, denying SB 79 projects near OTC in "
               "high-resource areas triggers an automatic Housing Accountability Act "
               "presumption. OTC sits squarely in a high-resource zone.")
    add_bullet(doc, "A CPRA request for internal correspondence on OTC's tier "
               "designation would surface whether staff computed an aggregate and "
               "withheld it, or never ran one. Either answer is damaging.")

    add_heading(doc, "The Woodside Echo", level=1)
    doc.add_paragraph(
        "Woodside declared itself a mountain-lion habitat in 2022 to dodge SB 9 "
        "duplexes. National ridicule followed. Reversal took two weeks."
    )
    doc.add_paragraph(
        "Oceanside's move is the same species: an arithmetic falsehood wearing "
        "planning-judgment clothing. 130 > 72 takes no more expertise than recognizing "
        "that a wealthy suburb is not exclusively inhabited by cougars. One difference: "
        "Woodside folded. Oceanside doubled down."
    )

    add_callout(doc,
        "Staff had the numbers. They chose which ones to present. 130 runs stop at "
        "OTC every weekday, from four separate operators. The bar is 72. No interpretation "
        "is required. Only a decision not to count.")

    doc.save(os.path.join(OUT_DIR, "4 - The OTC Misclassification.docx"))
    print("Built: 4 - The OTC Misclassification.docx")


# ============================================================
# DOCUMENT 5: COMPARABLE CITIES
# ============================================================
def build_comparables():
    doc = Document()
    style_doc(doc)
    add_title(doc, "What Happened to Other Cities That Fought",
              "Comparable Outcomes for Context | June 2026")

    add_heading(doc, "The Universal Pattern", level=1)
    doc.add_paragraph(
        "Every California city that has fought state housing law has followed the same arc:"
    )
    add_callout(doc, "Resistance → Litigation → Escalating costs → Political fracture → Capitulation")
    doc.add_paragraph(
        "No city has won. Not one. The question is never whether the city complies. It's "
        "how much money it burns before complying."
    )

    add_heading(doc, "City-by-City", level=1)

    cities = [
        ("Huntington Beach: Open Defiance",
         "Refused Housing Element update. Sued the state. Passed ballot measure overriding state law.",
         "4+ years of litigation, $10–$15M+ spent. $50K/month fines escalating. Still fighting. "
         "Court threatened compliance receiver.",
         "The $50K/month fine structure. If Oceanside triggers AG enforcement, the same framework applies."),

        ("La Cañada Flintridge: Builder's Remedy Denial",
         "Denied a Builder's Remedy project. Fought in court. Appealed.",
         "~$3.3M total cost. City dropped appeal after court ordered a $14M appeal bond. Couldn't afford it.",
         "The appeal bond mechanism. If Oceanside loses at trial and appeals, a bond potentially "
         "exceeding the project's value will be required. A city with a slim surplus can't post $14M bonds."),

        ("Encinitas: Voter Override (25 years)",
         "Voters rejected Housing Element three times. City argued ballot measure overrode state law.",
         "25 years without a compliant Housing Element. Judge finally ordered compliance. Every dollar "
         "spent on resistance was wasted.",
         "The endurance cost. Encinitas fought for a quarter century and lost everything. "
         "Oceanside's strategy is more sophisticated but the endgame is identical."),

        ("Woodside: Mountain Lion Habitat",
         "Declared entire town a mountain lion habitat to avoid SB 9 duplexes.",
         "National humiliation. Reversed in 2 weeks.",
         "Oceanside's Tier misclassification (130 > 72) is this generation's mountain lion claim. "
         "An arithmetic falsehood dressed as a planning judgment. Difference: Oceanside is doubling down."),

        ("Redondo Beach: Overlay Evasion",
         "Used housing overlay zones on commercial land to claim RHNA compliance without actually rezoning.",
         "Court of Appeal invalidated Housing Element. Supreme Court denied review. Builder's Remedy reactivated.",
         "Paper compliance doesn't work. Oceanside's SB 79 \"phasing ordinance\" is the same play."),
    ]

    for title, strategy, outcome, lesson in cities:
        add_heading(doc, title, level=2)
        add_bold_para(doc, "Strategy: ", strategy)
        add_bold_para(doc, "Outcome: ", outcome)
        add_bold_para(doc, "Lesson for Oceanside: ", lesson)
        doc.add_paragraph()

    add_heading(doc, "Projected Oceanside Outcome", level=1)
    doc.add_paragraph(
        "Most similar to Encinitas. Sophisticated, coastal, multi-year resistance that "
        "eventually collapses under court order. Estimated 2–4 years. Key difference: "
        "Oceanside has a Tier 1 transit station making the legal case far simpler (130 > 72) "
        "than Encinitas's complex RHNA allocation fights."
    )

    add_heading(doc, "Probability Scenarios", level=1)
    add_table(doc,
              ["Scenario", "Probability", "Timeline", "Cost"],
              [
                  ["Early capitulation (new majority or AG panic)", "20%",
                   "Late 2027", "$3M–$8M"],
                  ["Fight-then-fold (trial loss + rising fines)", "55%",
                   "Mid-2028", "$15M–$45M"],
                  ["Full resistance (through RHNA deadline)", "20%",
                   "2029–2030", "$80M–$249M"],
                  ["Coastal Act blocks SB 79 (pyrrhic victory)", "5%",
                   "2028–2029", "$10M–$20M"],
              ])

    doc.add_paragraph(
        "Even in the 5% scenario where the Coastal Act blocks SB 79, the legislature passes "
        "a trailer bill within 12 months. The city would spend $20M to buy one year of obstruction."
    )

    doc.save(os.path.join(OUT_DIR, "6 - What Happened to Other Cities.docx"))
    print("Built: 6 - What Happened to Other Cities.docx")


# ============================================================
# DOCUMENT 5: THE EXEMPTION STRATEGY (walking paths + full ordinance)
# ============================================================
def build_exemption_strategy():
    doc = Document()
    style_doc(doc)
    add_title(doc, "How Oceanside Exempted an Entire City",
              "Walking Paths, Deferments, and the SB 79 Evasion Ordinance | June 2026")

    add_heading(doc, "The Scale of What They're Attempting", level=1)
    doc.add_paragraph(
        "SANDAG parcel data shows approximately 18,800 parcels within half a mile of "
        "Oceanside's 8 transit stations. That is the SB 79 universe. On June 3, 2026, "
        "Oceanside's City Council took up an ordinance (File #26-1529) designed to exempt, "
        "exclude, and defer as many of those parcels as possible. Staff didn't pretend "
        "otherwise. Direct quote from the report: \"exempting and deferring as many sites "
        "as possible would potentially reduce SB 79's impact to the greatest extent possible.\" "
        "That's not our interpretation. Those are their words."
    )
    doc.add_paragraph(
        "The ordinance deploys multiple tools. The only one with verifiable parcel-level "
        "data — the walking path exemption — covers 710 parcels, or 3.8% of the total. "
        "The city claims deferments, rent control reclassification, habitat exclusions, "
        "and a planned Alternative Plan cover the rest. No public data confirms this."
    )

    add_heading(doc, "Prong 1: Walking Path Exemptions (Permanent)", level=1)
    doc.add_paragraph(
        "SB 79 has a narrow escape valve. If a parcel has no \"walking path\" of less than "
        "one mile to a transit stop, the city can exempt it (Gov. Code 65912.160(e)(1)). "
        "Requires site-specific findings backed by substantial evidence. The Legislature meant "
        "this for parcels cut off by freeways, rivers, rail corridors. Genuinely stranded land."
    )
    doc.add_paragraph(
        "Oceanside turned that escape valve into a fire hose."
    )

    add_heading(doc, "The Definition Trick", level=2)
    doc.add_paragraph(
        "The statute says \"walking path.\" Two words. No definition provided. So the city "
        "wrote one:"
    )
    add_callout(doc,
        "\"Walking path\" for this purpose means: a publicly accessible, continuous, and "
        "unobstructed path of travel with continuous paved or improved sidewalk from the "
        "closest point of the affected parcel to the pedestrian access point of the "
        "transit-oriented development stop.")
    doc.add_paragraph(
        "Look at the word \"continuous.\" One missing sidewalk segment anywhere along the "
        "route and the whole path fails. Doesn't matter if 95% of the route has perfect "
        "concrete. One gap. Chain broken. Parcel exempted."
    )
    doc.add_paragraph(
        "Staff knew exactly what they built. At the May 18 Planning Commission hearing: "
        "\"There is no definition of walking path in the statute, so the city created our "
        "own definition.\" They said it out loud."
    )

    add_heading(doc, "The Data That Contradicts the Conclusion", level=2)
    doc.add_paragraph(
        "A PRA request pulled two shapefiles from city GIS. Side by side, they blow the "
        "whole thing apart."
    )
    add_table(doc,
              ["Shapefile", "Contents", "Result"],
              [
                  ["pathoftravel.shp", "710 parcels analyzed for walkability",
                   "ALL 710 scored Walkability = 0"],
                  ["OceansideSB79_sidewalks.shp", "8,128 road segments with sidewalk status",
                   "1,726 in study area. 71% have sidewalks"],
              ])

    add_heading(doc, "Sidewalk Inventory Breakdown (Study Area)", level=3)
    add_table(doc,
              ["Status", "Segments", "Percentage"],
              [
                  ["Has sidewalk", "1,229", "71%"],
                  ["No sidewalk", "315", "18%"],
                  ["Other", "182", "11%"],
                  ["Total evaluated", "1,726", "100%"],
              ])

    add_heading(doc, "Where the 710 Parcels Actually Are — and Where They Aren't", level=3)
    doc.add_paragraph(
        "SANDAG parcel data reveals the true scale of what the walking path tool does and "
        "doesn't cover. Cross-referencing the PRA shapefile against the full parcel inventory:"
    )
    add_table(doc,
              ["Station", "Total Parcels (0.5 mi)", "Walking Path Shapefile", "Coverage"],
              [
                  ["OTC (130+ trains)", "4,572", "2", "0.04%"],
                  ["Coast Highway", "5,297", "53", "1.0%"],
                  ["Crouch Street", "1,334", "511", "38%"],
                  ["Civic Center", "974", "89", "9%"],
                  ["El Camino Real", "562", "40", "7%"],
                  ["Rancho Del Oro", "2,220", "3", "0.1%"],
                  ["College Blvd", "1,813", "11", "0.6%"],
                  ["Melrose", "2,031", "2", "0.1%"],
                  ["TOTAL", "~18,800", "710", "3.8%"],
              ])
    doc.add_paragraph(
        "The walking path exemption is not the main tool. It's a sideshow. It covers 3.8% "
        "of the SB 79 parcel universe, concentrated in the inland Sprinter corridor. OTC — "
        "the city's biggest station at 130+ daily trains with 4,572 parcels in its zone — has "
        "2 parcels in the walking path dataset. Two."
    )
    doc.add_paragraph(
        "The city claims deferments, rent control reclassification, and habitat exclusions "
        "cover the rest. But no parcel-level data for those tools has been produced. The gap "
        "between documented exemptions (3.8%) and the city's stated intent (\"to the greatest "
        "extent possible\") is itself the story. Either the other tools have massive holes and "
        "SB 79 will apply to thousands of parcels on July 1, or the city has parcel-level "
        "exemption data it hasn't disclosed. A PRA request for the deferment parcel list would "
        "resolve the question."
    )

    doc.add_paragraph(
        "Sit with those two datasets for a second. The city's own sidewalk inventory says "
        "71% of road segments near stations have sidewalks. Seven out of ten. And yet their "
        "walkability model scored every single parcel as unreachable. 710 tested, 710 failed. "
        "People walk to those stations every day. The data says they can't."
    )
    doc.add_paragraph(
        "This isn't a sidewalk coverage problem. It's a rigged test. Eighteen percent of "
        "segments lack sidewalks, scattered across the network in patches. Under the "
        "\"continuous\" standard, one gap anywhere along any possible route kills the entire "
        "path. Thirty years of skipping pedestrian infrastructure means there's always a "
        "gap somewhere. And now that neglect has a second job: blocking housing."
    )

    add_heading(doc, "What 'Walking Path' Actually Means", level=2)
    doc.add_paragraph(
        "The statute says \"walking path.\" Not \"continuous improved sidewalk.\" CalHDF's "
        "May 20 letter spells out what existing law already establishes:"
    )
    add_bullet(doc, "A walking path is any route traveled on foot. Streets, trails, paseos, "
               "pedestrian bridges, stairways. Sidewalks are one option, not a prerequisite.")
    add_bullet(doc, "Vehicle Code Section 21954 already says pedestrians can enter roadways "
               "without sidewalks. Yield to vehicles, exercise due care. That's it. Legal right "
               "to walk. No concrete required.")
    add_bullet(doc, "AB 2147 (the jaywalking reform bill, 2022) drove this home. The author's "
               "statement: \"People who need to walk in their neighborhoods should not be penalized "
               "for decades of infrastructure neglect and auto-first street design that fails to "
               "consider the needs of users who aren't in cars.\"")
    doc.add_paragraph(
        "The Legislature looked at exactly this scenario, where sidewalks don't exist, and "
        "said: people still walk there and that walking counts. Oceanside's ordinance flips "
        "it. No concrete, no path, no housing."
    )

    add_heading(doc, "Specific Parcels Wrongly Exempted", level=2)
    doc.add_paragraph(
        "CalHDF named specific neighborhoods where residents obviously walk to transit but "
        "the city's definition kicked them out:"
    )
    add_bullet(doc, "Moon Valley nurseries property west of Rancho Del Oro Station")
    add_bullet(doc, "Parcels immediately south of El Camino Real Station (some abut the station itself)")
    add_bullet(doc, "Skyline Drive and Sonja Court, southeast of El Camino Real Station")
    add_bullet(doc, "Neighborhoods southeast of Crouch Street Station, accessible via "
               "Crouch Street, Downs Street, Grandview Street")
    add_bullet(doc, "Neighborhoods northeast of Crouch Street Station, accessible via "
               "Crouch Street, Hoover Street")
    add_bullet(doc, "Parcels along S. Nevada Street, north of Coast Highway Station")
    doc.add_paragraph(
        "Some of these properties physically abut station land. You can stand on the parcel "
        "and see the platform. Exempted. Because somewhere between point A and point B, a "
        "chunk of sidewalk is missing."
    )

    add_heading(doc, "Prong 2: Deferment to 2032", level=1)
    doc.add_paragraph(
        "Everything the walking path exemption missed gets frozen. SB 79 lets cities defer "
        "implementation on qualifying sites until one year after adopting the seventh Housing "
        "Element revision (Gov. Code 65912.161(b)(1)). For Oceanside, that deadline stretches "
        "to June 15, 2032. Six full years before the law actually touches a single deferred parcel."
    )
    doc.add_paragraph(
        "Staff recommended deferring every site that fits any category:")

    add_table(doc,
              ["Deferment Category", "Code Section", "What It Covers"],
              [
                  ["Existing density at 50%+ of SB 79",
                   "65912.161(b)(1)(A)",
                   "Sites already zoned at half of SB 79 standards. Includes Downtown, Coastal Zone, Oceanside Blvd corridor."],
                  ["Low-resource area, 40%+ aggregate density",
                   "65912.161(b)(1)(B)(ii)",
                   "TOD zones primarily in low-resource areas per CTCAC maps. Only OTC zone meets the 40% threshold."],
                  ["Historic resources",
                   "65912.161(b)(1)(F)",
                   "7 properties in Townsite Neighborhood Planning Area on local register."],
                  ["Fire hazard / sea level rise",
                   "65912.161(b)(1)(D)-(E)",
                   "None found as of April 2026, but ordinance reserves authority to add them."],
              ])

    doc.add_paragraph(
        "Staff told council to layer both tools on every eligible site. Walking path "
        "exemption plus deferment, stacked. If one gets struck down, the other still blocks. "
        "Belt and suspenders."
    )

    add_heading(doc, "The Low-Resource Exemption Falls Apart at Tier 1", level=2)
    doc.add_paragraph(
        "This is where the train count fraud (see companion document on OTC misclassification) "
        "makes the deferment strategy collapse. The low-resource deferment under "
        "65912.161(b)(1)(B)(ii) requires local zoning to \"cumulatively allow for at least "
        "40 percent of the aggregate density\" under SB 79 for that TOD zone."
    )
    doc.add_paragraph(
        "Around OTC, current zoning allows 7 to 43 du/acre. Tier 1 stations get 100-120 "
        "du/acre within a quarter mile. Forty percent of 120 is 48. Oceanside maxes out at "
        "43. Falls short. The arithmetic only clears the bar if you pretend OTC is Tier 2 "
        "(80-100 du/acre), where 43 squeaks past at 43-54% of max. Fix the tier and the "
        "entire low-resource deferment for the OTC zone vanishes."
    )

    add_heading(doc, "Prong 3: The Alternative Plan (Delay by Design)", level=1)
    doc.add_paragraph(
        "SB 79 offers cities a legitimate planning tool: a TOD Alternative Plan "
        "(65912.161(a)) that redistributes capacity across sites. Concentrate density on "
        "commercial corridors, pull it back near single-family streets, keep the total "
        "unchanged. Real planning for real tradeoffs. Oceanside says it's working on one. "
        "Won't be ready by July 1."
    )
    doc.add_paragraph(
        "Which is the whole point. Staff told council the plan \"will require additional "
        "staff time\" and projected completion \"by late 2026 or early 2027.\" Meanwhile "
        "the deferment ordinance locks everything down through 2032. The alt plan provides "
        "the justification for a six-year freeze, whether it ever gets adopted or not."
    )
    doc.add_paragraph(
        "If the city finishes the plan, it can selectively unfreeze individual sites. Or it "
        "can leave them all frozen. The ordinance preserves that choice. Maximum flexibility, "
        "minimum housing."
    )

    add_heading(doc, "The Rent Control Overreach", level=1)
    doc.add_paragraph(
        "SB 79 draws a line on demolition: you can't tear down more than two units under "
        "rent or price control occupied in the past seven years (65912.157(h)(1)). Sensible "
        "protection for existing tenants. Nobody disputes the intent."
    )
    doc.add_paragraph(
        "But Oceanside stretched that protection to snap three categories it was never "
        "meant to cover:"
    )
    add_bullet(doc, "Sites subject to the Tenant Protection Act of 2019 (Civ. Code 1947.12). "
               "This is Good Cause eviction protection, not rent control. The TPA is in Division 3, "
               "Part 4, Title 5, Chapter 2 of the Civil Code. Rent control is in Chapter 2.7. "
               "Different chapters, different legal regimes. Costa-Hawkins Section 1954.52 "
               "explicitly allows property owners to set initial rents on post-1995 buildings. "
               "The Supreme Court held in Briggs v. Eden Council (1999) that different statutory "
               "terms carry different meanings.",
               "Good Cause eviction: ")
    add_bullet(doc, "Mobile home parks subject to Chapter 16B of the Oceanside City Code. "
               "These are local rent ceiling adjustments, but the city swept them into SB 79's "
               "\"price control\" exclusion without distinguishing which sites had active tenants "
               "in the past seven years.",
               "Mobile home parks: ")
    add_bullet(doc, "Sites with deed-restricted affordable housing. These units are restricted "
               "via subsidy, not via police power. The statute covers rent or price control "
               "\"through a public entity's valid exercise of its police powers.\" Deed "
               "restrictions through housing programs are not police power actions.",
               "Deed-restricted units: ")
    doc.add_paragraph(
        "CalHDF's assessment: \"impermissibly broad.\" The city grabbed tenant protection "
        "language meant to prevent displacement and repurposed it as a development ban on "
        "land the Legislature never intended to freeze."
    )

    add_heading(doc, "The Habitat and Wetland Exclusions", level=1)
    doc.add_paragraph(
        "Staff mapped sites within habitat preserves or 100 feet of wetlands per General "
        "Plan Policy 5-12, labeled them \"effectively undevelopable,\" and excluded them from "
        "SB 79 coverage."
    )
    doc.add_paragraph(
        "Problem: SB 79 has no habitat exception. Zero. The state law lists its own "
        "environmental carve-outs (fire hazard zones, sea level rise). Policy 5-12 is a "
        "local document. The California Constitution settles which one wins. Article XI, "
        "Section 7: local governments \"cannot choose to forgo application of land use "
        "policies imposed by the Legislature.\" Ailanto Properties, 142 Cal.App.4th at 595. "
        "State law eats local policy for breakfast."
    )
    doc.add_paragraph(
        "CalHDF put a finer point on it: the ordinance \"states in the staff report and "
        "maps, in effect, that state land use laws (i.e. SB 79) do not apply to certain "
        "areas of the City based on local ordinance and the local general plan.\" "
        "That's preemption in reverse. It doesn't work."
    )

    add_heading(doc, "The Procedural Defect", level=1)
    doc.add_paragraph(
        "Zoning ordinances follow a two-step process (Gov. Code 65854 et seq.). Planning "
        "Commission holds a hearing, votes, issues a recommendation. City Council notices "
        "its own hearing after that recommendation is published. After. Not before."
    )
    doc.add_paragraph(
        "Oceanside noticed the council hearing before the Planning Commission even voted on "
        "May 18. Two days later, the recommendation still wasn't posted. The vote field in "
        "the ordinance text was blank. Environmental Defense Project v. County of Sierra "
        "(2008, 158 Cal.App.4th 877): notice has to come after the commission acts. The "
        "purpose is to give the public a chance to respond to the actual recommendation, "
        "not staff's prediction of what it might be. The rushed timeline \"deprived the "
        "public of the opportunity to respond\" (Gov. Code 65033)."
    )
    doc.add_paragraph(
        "This alone could void the ordinance. A court wouldn't need to reach any of the "
        "substantive arguments."
    )

    add_heading(doc, "The Exhibit Mismatch", level=2)
    doc.add_paragraph(
        "Smaller defect, same root cause. The ordinance cites Exhibits A through F (lettered). "
        "The staff report attaches documents numbered 1 through 12. CalHDF caught it: \"the "
        "ordinance refers to maps (e.g. Exhibits A, B, D, E, and F). However, the attachments "
        "to the staff report are numbered, not lettered, and it is unclear which maps the "
        "ordinance is referring to.\""
    )
    doc.add_paragraph(
        "When your ordinance can't match its own maps to its own attachments, you wrote it "
        "in a hurry. SB 79 hits July 1. Staff was racing the clock."
    )

    add_heading(doc, "What This Adds Up To", level=1)
    doc.add_paragraph(
        "The city deployed multiple tools against ~18,800 SB 79 eligible parcels. "
        "Here's what we can verify and what we can't:"
    )
    add_bullet(doc, "Walking path exemptions: 710 parcels tested, 710 failed (100% rate). "
               "But that's 3.8% of the total SB 79 universe. Concentrated around Crouch Street "
               "and Civic Center. The definition is fabricated — the city's own sidewalk data "
               "contradicts it — but even if it holds, it barely dents the total.",
               "Verified: ")
    add_bullet(doc, "Deferments are claimed for parcels near all stations through 2032. "
               "The qualifying categories (existing density, low-resource, historic, fire, sea "
               "level rise) have specific requirements. No parcel-level data has been produced "
               "showing which parcels actually qualify. The low-resource exemption fails at Tier 1 "
               "densities around OTC.",
               "Claimed but unverified: ")
    add_bullet(doc, "Good Cause eviction and habitat exclusions cover additional parcels, "
               "but Good Cause only applies to existing rental housing 15+ years old, and habitat "
               "is limited to parcels near mapped wetlands and preserves.",
               "Claimed but unverified: ")
    add_bullet(doc, "The Alternative Plan is not ready and may never arrive. "
               "Staff projected late 2026 or early 2027. It provides justification for the "
               "deferment freeze whether it's ever adopted or not.",
               "Future: ")
    add_bullet(doc, "A procedural defect — the council hearing was noticed before the Planning "
               "Commission voted — could void the whole ordinance independently.",
               "Independently fatal: ")
    doc.add_paragraph(
        "Staff said the quiet part aloud. Reduce SB 79's impact \"to the greatest extent "
        "possible.\" Not \"implement thoughtfully.\" Not \"balance competing interests.\" "
        "Minimize. Everything. Period."
    )

    add_heading(doc, "The HCD Clock", level=1)
    doc.add_paragraph(
        "HCD received the draft ordinance May 19, 2026. Confirmed receipt May 26. After "
        "adoption, the city gets 60 days to submit the final version. If HCD objects, two "
        "options: amend the ordinance, or adopt a resolution arguing it already complies. "
        "If the city stonewalls, HCD kicks it to the Attorney General."
    )
    doc.add_paragraph(
        "CalHDF and Californians for Homeownership aren't sitting around waiting for the "
        "administrative process to grind through. Their May 20 letter already documents six "
        "independent legal defects. The litigation track started running before the council "
        "even voted."
    )

    add_callout(doc,
        "~18,800 parcels in Oceanside's SB 79 zones. The walking path exemption — the only "
        "tool with verifiable parcel data — covers 710 of them. 3.8%. The city says deferments "
        "and other tools cover the rest, but has produced no data proving it. Staff said the "
        "quiet part aloud: reduce impact \"to the greatest extent possible.\" Whether they "
        "succeeded depends on data they haven't disclosed. A PRA request for the deferment "
        "parcel list is the next move.")

    doc.save(os.path.join(OUT_DIR, "5 - The Exemption Strategy.docx"))
    print("Built: 5 - The Exemption Strategy.docx")


if __name__ == "__main__":
    build_summary()
    build_legal_strategy()
    build_costs()
    build_otc_misclassification()
    build_exemption_strategy()
    build_comparables()
    print(f"\nAll documents saved to: {OUT_DIR}")
