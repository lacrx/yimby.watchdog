#!/usr/bin/env python3
"""Generate Word document: Rebuttal to density cap deflection arguments."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    hs.font.name = 'Calibri'

# ─── Title ───
title = doc.add_heading(
    'Downtown Oceanside Filing Collapse\nRebuttal to Alternative Explanations', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Data Analysis — June 2026')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# ═══════════════════════════════════════════
# Executive Summary
# ═══════════════════════════════════════════

doc.add_heading('Executive Summary', level=2)

doc.add_paragraph(
    'The data doesn\'t support the interpretation that construction costs or land '
    'availability caused the downtown filing collapse. Rest-of-city Oceanside filings went '
    'up 49% during the exact same period downtown dropped 96%. If construction costs were '
    'the cause, both would decline — they face identical lumber, concrete, and interest rates. '
    'San Diego County hit a 17-year permit record during this period.'
)

doc.add_paragraph(
    'Developers respond to regulatory signals at adoption, not certification. The filing '
    'collapse tracks to October 2023 (council adoption of the 86 du/acre cap), not February '
    '2026 (CCC certification). That\'s how capital markets work.'
)

doc.add_paragraph(
    'And "all the major parcels have approvals" — most of those approvals haven\'t broken '
    'ground. The 9-block area isn\'t the only developable downtown land. The density cap makes '
    'remaining sites infeasible, which is exactly what a 96% collapse shows.'
)

doc.add_paragraph(
    'Three alternative explanations have been offered for the 96% decline in downtown '
    'Oceanside housing filings from 2022 to 2025. Each is testable against the data. '
    'All three fail.'
)

# ═══════════════════════════════════════════
# SECTION 1: The Data
# ═══════════════════════════════════════════

doc.add_heading('The Filing Data', level=2)

doc.add_paragraph(
    'Source: HCD Annual Progress Report Table A (data.ca.gov). Downtown defined by '
    'D-District bounding box (lat 33.1858–33.2096, lon -117.3923–-117.3749). '
    'All figures are proposed units filed by application date.'
)

table = doc.add_table(rows=9, cols=5, style='Light Shading Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Year', 'Downtown', 'Rest of City', 'Total', 'Downtown %']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

data = [
    ('2018', '76', '500', '576', '13%'),
    ('2019', '416', '674', '1,090', '38%'),
    ('2020', '20', '379', '399', '5%'),
    ('2021*', '471', '1,227', '1,698', '28%'),
    ('2022', '1,946', '1,265', '3,211', '61%'),
    ('2023**', '1,280', '697', '1,977', '65%'),
    ('2024', '816', '1,207', '2,023', '40%'),
    ('2025', '69', '1,885', '1,954', '4%'),
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph(
    '* 43 du/acre cap removed November 2021.  '
    '** 86 du/acre cap adopted October 2023.'
).italic = True

doc.add_paragraph(
    'Downtown filings fell from 1,946 units (2022) to 69 units (2025) — a 96% decline. '
    'In the same period, rest-of-city filings rose from 1,265 to 1,885 — a 49% increase.'
)

# ═══════════════════════════════════════════
# SECTION 2: Claim 1 — Construction Costs
# ═══════════════════════════════════════════

doc.add_heading('Claim 1: "Construction costs and borrowing costs"', level=2)

doc.add_heading('The argument', level=3)
doc.add_paragraph(
    'The slowdown is primarily a result of construction costs (both material and the '
    'cost of borrowing).'
)

doc.add_heading('Why it fails', level=3)
doc.add_paragraph(
    'Construction costs and interest rates are market-wide forces. They hit every project '
    'in every neighborhood equally. If they caused the downtown collapse, rest-of-city '
    'filings would decline too. They did the opposite.'
)

p = doc.add_paragraph()
p.add_run('The control group: ').bold = True
p.add_run(
    'Rest-of-city Oceanside filed 1,582 units in multifamily projects (10+ units) in 2025 '
    'alone, including Ocean Kamp Residences (574 units), Tierra Norte Parcel B (229 units), '
    'and Olive Park Apartments (199 units). Same city, same interest rates, same lumber prices.'
)

p = doc.add_paragraph()
p.add_run('Regional context: ').bold = True
p.add_run(
    'San Diego County issued 11,673 housing permits in 2023 — a 17-year record, up 21% from '
    '2022, with 9,100 multifamily. The City of San Diego permitted 8,782 new homes in 2024. '
    'The regional market absorbed cost increases and kept building.'
)

p = doc.add_paragraph()
p.add_run('National context: ').bold = True
p.add_run(
    'National multifamily starts declined approximately 25% from the 2022 peak. A 25% market '
    'correction is normal. A 96% localized collapse in one neighborhood is not a market signal — '
    'it is a regulatory signal.'
)

p = doc.add_paragraph()
p.add_run('Cost stabilization: ').bold = True
p.add_run(
    'Construction cost inflation peaked at ~17% YOY in 2022 (ENR/Construction Analytics), '
    'then decelerated to 3.4% by 2024 — back near historical norms. Turner Building Cost '
    'Index showed 3.6% YOY growth in Q3 2024. Costs were stabilizing during the period '
    'downtown filings collapsed.'
)

# ═══════════════════════════════════════════
# SECTION 3: Claim 2 — Cap Timing
# ═══════════════════════════════════════════

doc.add_heading('Claim 2: "The density cap was literally just put in place last month"', level=2)

doc.add_heading('The argument', level=3)
doc.add_paragraph(
    'The cap had to go through the Coastal Commission and was only certified in February 2026. '
    'It could not have affected applications in 2024 or 2025.'
)

doc.add_heading('Why it fails', level=3)
doc.add_paragraph(
    'Developers and capital markets respond to regulatory signals at adoption, not at final '
    'legal certification. The council adopted the 86 du/acre cap in October 2023. From that '
    'point forward, any developer running a pro forma for a downtown site had to pencil at '
    '86 du/acre maximum.'
)

p = doc.add_paragraph()
p.add_run('Economic theory: ').bold = True
p.add_run(
    'Titman (1985, American Economic Review, 700+ citations) established that regulatory '
    'uncertainty increases the option value of holding land and rationally delays development. '
    'Developers do not need a law to take effect — the signal that it will is sufficient.'
)

p = doc.add_paragraph()
p.add_run('Portland precedent: ').bold = True
p.add_run(
    'Portland adopted mandatory inclusionary zoning in December 2016, effective February 2017. '
    'Before the effective date, developers rushed to file. After: zero private apartment '
    'projects of 20+ units were submitted for land use review. Annual permits fell from '
    '~6,000 to ~1,500. The market fully priced in the policy at adoption.'
)

p = doc.add_paragraph()
p.add_run('Lender behavior: ').bold = True
p.add_run(
    'Construction lenders require confirmed entitlements before closing. An adopted density '
    'cap that threatens unit counts directly undermines debt-service coverage ratios. Per the '
    'OCC Comptroller\'s Handbook, zoning compliance is a closing condition for CRE lending. '
    'Lenders pause at adoption, not certification.'
)

p = doc.add_paragraph()
p.add_run('Coastal zone chill: ').bold = True
p.add_run(
    'Circulate San Diego\'s 2024 report documented a "chilling effect" where developers avoid '
    'the coastal zone entirely due to years of CCC delay and uncertain outcomes. The prospect '
    'of "years of delay with no certain outcome" (Jeannette Temple, CalMatters) stops projects '
    'from being proposed at all.'
)

p = doc.add_paragraph()
p.add_run('The data confirms it: ').bold = True
p.add_run(
    'Downtown filings declined smoothly from adoption: -34% in 2023 (year of adoption), '
    '-36% in 2024, -92% in 2025. This is a regulatory chill curve, not a sudden event at '
    'certification. The only 2026 downtown filing (DB26-00001, 142 units) used state density '
    'bonus law to bypass local zoning entirely — evidence that developers treated the local '
    'regulatory environment as hostile well before CCC certification.'
)

# ═══════════════════════════════════════════
# SECTION 4: Claim 3 — Land Exhaustion
# ═══════════════════════════════════════════

doc.add_heading('Claim 3: "All major parcels have had plan approvals — '
                'we don\'t have infinite downtown land"', level=2)

doc.add_heading('The argument', level=3)
doc.add_paragraph(
    'The 9-block master plan parcels are all approved. Downtown land is effectively exhausted.'
)

doc.add_heading('Why it fails', level=3)

p = doc.add_paragraph()
p.add_run('The 9-block plan is not "downtown": ').bold = True
p.add_run(
    'The Nine Block Master Plan (approved April 2000) covers nine blocks near the pier, '
    'bounded by The Strand, Cleveland Street, Seagaze Drive, and Civic Center Drive. It was '
    'a narrow redevelopment plan for visitor-serving uses and housing on specific parcels. '
    'It was never the universe of downtown developable land.'
)

p = doc.add_paragraph()
p.add_run('Major projects exist outside the 9 blocks: ').bold = True

table2 = doc.add_table(rows=5, cols=4, style='Light Shading Accent 1')
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Project', 'Units', 'Status', 'Location']):
    cell = table2.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

projects = [
    ('401 Mission', '326', 'Approved Oct 2025, not built', 'Outside 9-block'),
    ('901 Mission', '273', 'Approved Oct 2025, not built', 'Outside 9-block'),
    ('Jefferson Ocean Creek', '295', 'Under construction', 'Outside 9-block'),
    ('Blocks 5 & 20', '373', 'Approved Jan 2026, not built', 'Within 9-block'),
]
for row_idx, row_data in enumerate(projects):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph(
    '894 units on downtown parcels outside the 9-block area, plus 373 approved within it. '
    '"Approved" does not mean "built" — most have not broken ground. The land is not '
    'exhausted. The projects are stuck.'
)

p = doc.add_paragraph()
p.add_run('The 9-block plan\'s actual housing delivery: ').bold = True
p.add_run(
    'The Nine Block Master Plan was adopted in April 2000. In 26 years, it has delivered '
    'approximately 176 completed housing units (Pierside North: 66, Pierside South: 110). '
    'Several blocks went to hotels (Mission Pacific, Seabird Resort) and condos (The SALT) '
    'rather than rental housing. Blocks 5 & 20 (373 units) were approved in January 2026 but '
    'have not broken ground. Saying "all major parcels have had plan approvals" conflates '
    'entitlement with delivery. A plan that produced ~176 housing units across 26 years is '
    'not evidence that downtown land is exhausted — it is evidence that the 9-block framework '
    'itself underperformed. The remaining downtown parcels outside the 9 blocks need more '
    'density to compensate for that underperformance, not an 86 du/acre cap.'
)

p = doc.add_paragraph()
p.add_run('The density cap is the binding constraint: ').bold = True
p.add_run(
    'If remaining downtown parcels are capped at 86 du/acre, fewer units can be built per '
    'site, which makes projects financially infeasible in a high-cost coastal market. '
    'This is exactly what a 96% filing collapse shows — not land exhaustion, but regulatory '
    'suppression of per-parcel yield.'
)

p = doc.add_paragraph()
p.add_run('SSCSP dwarfs downtown capacity: ').bold = True
p.add_run(
    'The Smart & Sustainable Corridors Specific Plan (SSCSP) covers Mission Avenue, '
    'Oceanside Boulevard, and Vista Way corridors with 17,000–20,000 units of new capacity '
    'on parcels currently zoned at 3.9 du/acre. The first filing (DB26-00001, 142 units at '
    '1640 Oceanside Blvd) arrived before the plan was even adopted. Downtown is a fraction '
    'of the city\'s development capacity.'
)

# ═══════════════════════════════════════════
# SECTION 5: Conclusion
# ═══════════════════════════════════════════

doc.add_heading('Conclusion', level=2)

doc.add_paragraph(
    'All three alternative explanations fail the same test: they predict effects that should '
    'appear city-wide or regionally, but the collapse is confined to downtown Oceanside. '
    'The only variable that changed exclusively for downtown was the density cap adopted in '
    'October 2023.'
)

table3 = doc.add_table(rows=5, cols=3, style='Light Shading Accent 1')
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Claim', 'Prediction', 'Data Shows']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    cell.paragraphs[0].runs[0].bold = True

conclusions = [
    ('Construction costs', 'City-wide decline', 'Rest of city UP 49%'),
    ('Borrowing costs', 'Regional decline', 'SD County hit 17-year record'),
    ('Cap not yet certified', 'No effect until Feb 2026',
     'Smooth decline from Oct 2023 adoption'),
    ('Land exhausted', 'No parcels available',
     '894 units approved outside 9-block, most unbuilt'),
]
for row_idx, row_data in enumerate(conclusions):
    for col_idx, val in enumerate(row_data):
        table3.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run(
    'Downtown Oceanside\'s housing filing collapse is a regulatory outcome, not a market one.')
run.bold = True

# ─── Sources ───

doc.add_heading('Sources', level=2)

sources = [
    'HCD Annual Progress Report Table A — data.ca.gov (2018–2025, JURIS_NAME = OCEANSIDE)',
    'eTRAKiT permit system — etrakit.cityofoceanside.com (2026 filings)',
    'Titman, S. (1985). "Urban Land Prices Under Uncertainty." American Economic Review.',
    'City Observatory — Portland inclusionary zoning impact (Joe Cortright)',
    'OCC Comptroller\'s Handbook — CRE lending and zoning compliance requirements',
    'Circulate San Diego (2024) — Coastal zone chilling effect on affordable housing',
    'Turner Building Cost Index — turnerconstruction.com (Q3 2024)',
    'ENR/Construction Analytics — construction cost inflation 2020–2025',
    'Times of San Diego — "SD County issued 11,673 housing permits" (July 2024)',
    'Inside San Diego — "City of San Diego permits nearly 8,800 new homes in 2024"',
    'Mainstreet Oceanside — Nine Block Master Plan history',
    'Coast News — Blocks 5 & 20 approval (Jan 2026), 401 Mission, 901 Mission',
    'Onward Oceanside — Smart & Sustainable Corridors Specific Plan capacity estimates',
]

for s in sources:
    doc.add_paragraph(s, style='List Bullet')

# ─── Save ───

out = '/home/thomas/repos/civics/share/oceanside-sb79-briefing/Downtown Filing Collapse Rebuttal.docx'
doc.save(out)
print(f'Saved: {out}')
